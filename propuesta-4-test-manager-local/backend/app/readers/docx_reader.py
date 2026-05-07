from __future__ import annotations
import io


class DocxReader:
    """Lee documentos Word (.docx), incluyendo tablas."""

    @property
    def supported_extensions(self) -> set[str]:
        return {".docx"}

    def read(self, content: bytes) -> str:
        from docx import Document
        doc = Document(io.BytesIO(content))

        parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())

        return "\n".join(parts)
